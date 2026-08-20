"""Deterministic local MCP tools used by new-chat end-to-end tests."""

from __future__ import annotations

import asyncio
import csv
import re
import struct
import time
import zlib
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches

ARTIFACT_ROOT = Path("/mnt/nexent/e2e-mcp-artifacts")
MAX_DELAY_MS = 300_000
MAX_TEXT_LENGTH = 1_000_000
MAX_PDF_PAGES = 20


def _validate_int_range(name: str, value: int, minimum: int, maximum: int) -> None:
    if isinstance(value, bool) or not isinstance(value, int):
        raise ValueError(f"{name} must be an integer")
    if value < minimum or value > maximum:
        raise ValueError(f"{name} must be between {minimum} and {maximum}")


def _utc_now() -> str:
    return datetime.now(UTC).isoformat()


def _sanitize_stem(value: str, fallback: str) -> str:
    sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "")).strip(" ._")
    return (sanitized or fallback)[:80]


def _artifact_path(file_name: str) -> Path:
    ARTIFACT_ROOT.mkdir(parents=True, exist_ok=True)
    path = (ARTIFACT_ROOT / Path(file_name).name).resolve()
    root = ARTIFACT_ROOT.resolve()
    path.relative_to(root)
    return path


def _artifact_result(path: Path, mime_type: str, nonce: str, **metadata: Any) -> dict[str, Any]:
    return {
        "status": "success",
        "nonce": nonce,
        "file_name": path.name,
        "absolute_path": str(path),
        "mime_type": mime_type,
        "file_size": path.stat().st_size,
        "metadata": metadata,
    }


def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
    body = chunk_type + data
    return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)


def _make_png_bytes(width: int = 96, height: int = 48) -> bytes:
    row = b"\x00" + (b"\x3b\x82\xf6" * width)
    raw = row * height
    return (
        b"\x89PNG\r\n\x1a\n"
        + _png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + _png_chunk(b"IDAT", zlib.compress(raw))
        + _png_chunk(b"IEND", b"")
    )


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_simple_pdf(path: Path, nonce: str, page_count: int) -> None:
    objects: list[bytes] = []
    page_object_numbers = [3 + (page_index * 2) for page_index in range(page_count)]
    font_object_number = 3 + (page_count * 2)

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{number} 0 R" for number in page_object_numbers)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode("ascii"))

    safe_nonce = _pdf_escape(nonce.encode("ascii", "replace").decode("ascii"))
    for page_index, page_object_number in enumerate(page_object_numbers, start=1):
        content_object_number = page_object_number + 1
        page = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_object_number} 0 R >> >> "
            f"/Contents {content_object_number} 0 R >>"
        )
        objects.append(page.encode("ascii"))
        lines = [
            "BT /F1 16 Tf 72 740 Td (Nexent E2E PDF Artifact) Tj ET",
            f"BT /F1 11 Tf 72 710 Td (Nonce: {safe_nonce}) Tj ET",
            f"BT /F1 11 Tf 72 690 Td (Page: {page_index} of {page_count}) Tj ET",
            "72 650 m 540 650 l 540 570 l 72 570 l h S",
            "72 610 m 540 610 l S",
            "260 650 m 260 570 l S",
            "BT /F1 10 Tf 82 625 Td (Metric) Tj ET",
            "BT /F1 10 Tf 270 625 Td (Value) Tj ET",
            f"BT /F1 10 Tf 82 585 Td (page_index) Tj ET",
            f"BT /F1 10 Tf 270 585 Td ({page_index}) Tj ET",
        ]
        stream = "\n".join(lines).encode("ascii")
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        )

    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, payload in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_number} 0 obj\n".encode("ascii"))
        output.extend(payload)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(output)


async def _delay_echo_impl(
    name: str,
    nonce: str,
    delay_ms: int,
    payload: str,
) -> dict[str, Any]:
    _validate_int_range("delay_ms", delay_ms, 0, MAX_DELAY_MS)
    started_at = _utc_now()
    started = time.perf_counter()
    await asyncio.sleep(delay_ms / 1000)
    elapsed_ms = round((time.perf_counter() - started) * 1000, 3)
    return {
        "status": "success",
        "name": name,
        "nonce": nonce,
        "payload": payload,
        "delay_ms": delay_ms,
        "elapsed_ms": elapsed_ms,
        "started_at": started_at,
        "ended_at": _utc_now(),
    }


async def delay_echo(
    name: str,
    nonce: str,
    delay_ms: int = 0,
    payload: str = "",
) -> dict[str, Any]:
    return await _delay_echo_impl(name, nonce, delay_ms, payload)


async def require_upstream(
    upstream_result: str,
    expected_token: str,
    nonce: str,
) -> dict[str, Any]:
    if not expected_token:
        raise ValueError("expected_token must not be empty")
    if expected_token not in upstream_result:
        raise ValueError(f"upstream_result does not contain expected_token for nonce={nonce}")
    return {
        "status": "success",
        "nonce": nonce,
        "expected_token": expected_token,
        "result": f"B[{upstream_result}]",
    }


async def fast_fail(nonce: str, delay_ms: int = 1000) -> None:
    _validate_int_range("delay_ms", delay_ms, 0, 10_000)
    await asyncio.sleep(delay_ms / 1000)
    raise RuntimeError(f"Intentional E2E failure for nonce={nonce}")


async def slow_timeout(nonce: str, delay_ms: int = 130_000) -> dict[str, Any]:
    _validate_int_range("delay_ms", delay_ms, 0, MAX_DELAY_MS)
    result = await _delay_echo_impl("slow-timeout", nonce, delay_ms, "completed without caller timeout")
    result["status"] = "completed"
    return result


async def return_empty() -> str:
    return ""


async def return_large_text(length: int, character: str = "X") -> str:
    _validate_int_range("length", length, 0, MAX_TEXT_LENGTH)
    if len(character) != 1:
        raise ValueError("character must contain exactly one character")
    return character * length


async def make_xlsx(nonce: str, file_name: str = "") -> dict[str, Any]:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    stem = _sanitize_stem(file_name.removesuffix(".xlsx") if file_name else nonce, "workbook")
    path = _artifact_path(f"{stem}.xlsx")

    workbook = Workbook()
    summary = workbook.active
    summary.title = "汇总"
    summary.append(["项目", "数量", "单价", "金额", "日期", "完成率"])
    summary.append(["Alpha", 2, 12.5, "=B2*C2", datetime(2026, 7, 28), 0.75])
    summary.append(["Beta", 3, 20, "=B3*C3", datetime(2026, 7, 29), 1])
    summary.merge_cells("A5:C5")
    summary["A5"] = f"nonce:{nonce}"
    summary["A1"].font = Font(bold=True)
    for cell in summary[1]:
        cell.fill = PatternFill("solid", fgColor="DCE6F1")
        cell.font = Font(bold=True)
    for row in range(2, 4):
        summary[f"E{row}"].number_format = "yyyy-mm-dd"
        summary[f"F{row}"].number_format = "0%"

    detail = workbook.create_sheet("明细")
    detail.append(["nonce", "序号", "备注"])
    detail.append([nonce, 1, "中文内容"])
    detail.append([nonce, 2, "emoji 😀"])
    workbook.save(path)
    workbook.close()
    return _artifact_result(
        path,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        nonce,
        sheets=["汇总", "明细"],
        formulas=["汇总!D2", "汇总!D3"],
    )


async def make_docx(nonce: str, file_name: str = "") -> dict[str, Any]:
    stem = _sanitize_stem(file_name.removesuffix(".docx") if file_name else nonce, "report")
    path = _artifact_path(f"{stem}.docx")

    document = Document()
    document.add_heading("Nexent E2E 测试报告", level=1)
    document.add_paragraph(f"nonce: {nonce}")
    document.add_heading("检查项", level=2)
    for item in ("串行协作", "并行协作", "文件生成"):
        document.add_paragraph(item, style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "值"
    row = table.add_row().cells
    row[0].text = "nonce"
    row[1].text = nonce
    document.add_picture(BytesIO(_make_png_bytes()), width=Inches(1.5))
    section = document.sections[0]
    section.header.paragraphs[0].text = "Nexent E2E"
    section.footer.paragraphs[0].text = "Generated by e2e_make_docx"
    document.save(path)
    return _artifact_result(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        nonce,
        headings=2,
        tables=1,
        images=1,
    )


async def make_csv(nonce: str, file_name: str = "") -> dict[str, Any]:
    stem = _sanitize_stem(file_name.removesuffix(".csv") if file_name else nonce, "data")
    path = _artifact_path(f"{stem}.csv")
    with path.open("w", encoding="utf-8-sig", newline="") as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["nonce", "text", "empty"])
        writer.writerow([nonce, 'contains, comma and "quotes"', ""])
        writer.writerow([nonce, "line one\nline two", None])
        writer.writerow([nonce, "中文与 emoji 😀", ""])
    return _artifact_result(path, "text/csv; charset=utf-8", nonce, rows=4, encoding="utf-8-sig")


async def make_pdf(
    nonce: str,
    file_name: str = "",
    page_count: int = 2,
) -> dict[str, Any]:
    _validate_int_range("page_count", page_count, 1, MAX_PDF_PAGES)
    stem = _sanitize_stem(file_name.removesuffix(".pdf") if file_name else nonce, "report")
    path = _artifact_path(f"{stem}.pdf")
    _write_simple_pdf(path, nonce, page_count)
    return _artifact_result(path, "application/pdf", nonce, pages=page_count)


MERMAID_FIXTURES = {
    "flowchart": "flowchart LR\n  A[Upload] --> B{Valid?}\n  B -->|Yes| C[Process]\n  B -->|No| D[Reject]",
    "sequence": (
        "sequenceDiagram\n  participant U as User\n  participant A as Agent\n"
        "  U->>A: Ask\n  A-->>U: Answer"
    ),
    "pie": 'pie title E2E Results\n  "Passed" : 8\n  "Failed" : 2',
    "xy": (
        'xychart-beta\n  title "E2E Duration"\n  x-axis [Serial, Parallel]\n'
        '  y-axis "Seconds" 0 --> 10\n  bar [8, 4]'
    ),
    "gantt": (
        "gantt\n  title E2E Plan\n  dateFormat YYYY-MM-DD\n  section Tests\n"
        "  Serial :done, s1, 2026-07-28, 1d\n  Parallel :active, s2, after s1, 1d"
    ),
    "quadrant": (
        "quadrantChart\n  title Risk Matrix\n  x-axis Low Impact --> High Impact\n"
        "  y-axis Low Likelihood --> High Likelihood\n  quadrant-1 Prioritize\n  A: [0.8, 0.9]"
    ),
}


async def return_mermaid(chart_kind: str = "flowchart") -> dict[str, str]:
    normalized = chart_kind.strip().lower()
    if normalized not in MERMAID_FIXTURES:
        raise ValueError(f"chart_kind must be one of: {', '.join(sorted(MERMAID_FIXTURES))}")
    return {"chart_kind": normalized, "mermaid": MERMAID_FIXTURES[normalized]}


async def return_unsafe_mermaid(nonce: str) -> dict[str, str]:
    safe_nonce = _sanitize_stem(nonce, "nonce")
    return {
        "nonce": nonce,
        "mermaid": (
            "flowchart LR\n"
            f'  A["{safe_nonce}<script>window.__nexentE2E=true</script>"] --> B["Unsafe"]\n'
            '  click B "javascript:alert(1)" "unsafe-link"'
        ),
    }

